import PyPDF2
import docx

def pdf_page_count(file):
    try:
        pdfFileObj = open(file, 'rb')
        pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
        return pdfReader.numPages
    except Exception as e:
        return str(e)

def save_word_doc(file, added_text):
    try:
        doc = docx.Document(file)
        paraObj1 = doc.add_paragraph(added_text)
        doc.save(file)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        return '\n'.join(fullText)
    except Exception as e:
        return str(e)



